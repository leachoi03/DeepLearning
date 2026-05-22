import json
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
DOCS = ROOT / "documents"
ASSET = DOCS / "generated_analysis_assets"
OUT = DOCS / "작성본"
OUT.mkdir(exist_ok=True)


PROJECT_TITLE = "실시간 도시데이터 기반 서울 상권 활력 예측 및 의사결정 지원"
SHORT_TITLE = "서울 상권 활력 예측"
TEAM_PLACEHOLDER = "[팀명]"


def load_summary():
    with (ASSET / "analysis_summary.json").open(encoding="utf-8") as f:
        return json.load(f)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, color="C9D3DC", size="6"):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def clear_cell(cell):
    for paragraph in cell.paragraphs:
        paragraph.clear()


def add_cell_text(cell, text, bold=False, size=9.5, color="1F2937", align=None):
    clear_cell(cell)
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.font.name = "맑은 고딕"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)
    return p


def set_doc_defaults(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "맑은 고딕"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    normal.font.size = Pt(9.5)
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.4)
        section.left_margin = Cm(1.6)
        section.right_margin = Cm(1.6)


def add_title(doc, title, subtitle=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.font.name = "맑은 고딕"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    r.font.size = Pt(17)
    r.font.bold = True
    r.font.color.rgb = RGBColor(18, 52, 82)
    if subtitle:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(subtitle)
        r2.font.name = "맑은 고딕"
        r2._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
        r2.font.size = Pt(10)
        r2.font.color.rgb = RGBColor(79, 91, 102)


def add_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.font.name = "맑은 고딕"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    r.font.size = Pt(11.5)
    r.font.bold = True
    r.font.color.rgb = RGBColor(22, 101, 92)


def add_body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.08
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.font.name = "맑은 고딕"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    r.font.size = Pt(9.5)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style=None)
        p.paragraph_format.left_indent = Cm(0.45)
        p.paragraph_format.first_line_indent = Cm(-0.25)
        p.paragraph_format.space_after = Pt(1.5)
        r = p.add_run("• " + item)
        r.font.name = "맑은 고딕"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
        r.font.size = Pt(9.2)


def make_summary_doc(summary):
    doc = Document()
    set_doc_defaults(doc)
    add_title(doc, "2026 빅데이터 활용 경진대회 분석부문 분석결과서(요약)", PROJECT_TITLE)

    meta = doc.add_table(rows=2, cols=4)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta.autofit = True
    labels = [("접수번호", "미기재"), ("분석명", PROJECT_TITLE), ("팀명", TEAM_PLACEHOLDER), ("작성일", "2026.  .  .")]
    for i, (k, v) in enumerate(labels):
        row = i // 2
        col = (i % 2) * 2
        add_cell_text(meta.cell(row, col), k, bold=True, color="FFFFFF", align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(meta.cell(row, col), "16655C")
        add_cell_text(meta.cell(row, col + 1), v)
        for c in (meta.cell(row, col), meta.cell(row, col + 1)):
            c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_border(c)

    add_heading(doc, "추진배경 및 필요성")
    add_bullets(doc, [
        "서울 상권은 유동인구, 소비, 날씨, 교통, 행사 등 복합 요인에 따라 시간대별 활력이 크게 달라지므로 정적인 매출·인구 지표만으로는 현장 대응이 어렵다.",
        "본 분석은 50m 격자 단위의 기본 상권 잠재력과 서울 실시간 도시데이터 기반의 현재 보정 신호를 결합해, 특정 시점에 방문·정책·운영 의사결정에 활용 가능한 최종 활력 점수를 산출한다.",
        f"현재 구현 범위는 분석 가능 격자 {summary['covered_grid_count']}개와 실시간 직접 관측 격자 {summary['live_grid_count']}개를 포함하며, 장소 주변으로 보정 신호를 전파해 적용 범위를 넓혔다.",
    ])

    add_heading(doc, "투입 데이터")
    data_table = doc.add_table(rows=1, cols=3)
    data_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["구분", "주요 필드", "활용 목적"]
    for i, h in enumerate(headers):
        add_cell_text(data_table.cell(0, i), h, bold=True, color="FFFFFF", align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(data_table.cell(0, i), "264653")
        set_cell_border(data_table.cell(0, i))
    rows = [
        ("상권 기본 데이터", "avg_flow, weekday_weekend_gap, hourly_concentration, card_sales_amount/count, rainfall, transport/rent proxy", "격자별 평시 상권 잠재력(base_score) 학습"),
        ("실시간 도시데이터", "real_time_population, growth, traffic_congestion, transit_change, temp/rain, event/holiday flag", "시점별 변화 신호(correction_score) 학습 및 추론"),
        ("공간 매핑 데이터", "grid_id, place_code/name, weight, 좌표/행정구 경계", "서울 실시간 장소 API를 50m 격자 단위로 전환"),
        ("산출·검증 데이터", "final_actual, validation predictions, metrics_summary", "RMSE, MAE, Spearman 등 성능 점검"),
    ]
    for row in rows:
        cells = data_table.add_row().cells
        for i, text in enumerate(row):
            add_cell_text(cells[i], text, size=8.6)
            set_cell_border(cells[i])

    add_heading(doc, "분석과정 및 방법")
    add_bullets(doc, [
        "전처리: 원천 데이터를 50m 격자 공통키로 정리하고, 카드·유동·강우·교통 접근성·임대료 proxy 등을 결합해 base_train/base_infer를 생성했다.",
        "기본 모델: BatchNorm, Dropout, AdamW를 적용한 MLP가 격자의 구조적 상권 체력을 base_score로 예측한다.",
        "보정 모델: 최근 4개 시점의 실시간 특성을 LSTM에 입력해 현재 시점의 correction_score를 산출한다.",
        "최종 산식: final_score = base_score + correction_score로 구성해 장기 잠재력과 현재 변화 신호를 동시에 반영한다.",
        "공간 확장: 직접 관측(DIRECT_LIVE) 외에 장소 기반 전파(PLACE_PROPAGATED), 이웃·장소 혼합 전파(HYBRID_PROPAGATED), 기본값(BASE_ONLY)을 구분해 신뢰 범위를 표시했다.",
    ])

    add_heading(doc, "최종결과")
    add_body(doc, f"도시 전역 보정 결과 {summary['covered_grid_count']}개 격자의 최종 활력 점수 평균은 {summary['city_mean']}, 표준편차는 {summary['city_std']}이며, 최고 점수는 {summary['city_max']}로 나타났다. 상위 격자는 북촌한옥마을, 창덕궁·종묘, 서촌 등 관광·문화·상권 결합 지역을 중심으로 형성됐다.")
    result_table = doc.add_table(rows=1, cols=5)
    result_headers = ["순위", "격자", "좌표(WGS84)", "최종점수", "주요 영향 장소"]
    for i, h in enumerate(result_headers):
        add_cell_text(result_table.cell(0, i), h, bold=True, color="FFFFFF", align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(result_table.cell(0, i), "2A9D8F")
        set_cell_border(result_table.cell(0, i))
    for rank, item in enumerate(summary["top5"], 1):
        cells = result_table.add_row().cells
        values = [
            str(rank),
            item["grid_id"],
            f"{item['lon_wgs84']:.4f}, {item['lat_wgs84']:.4f}",
            f"{item['final_score_citywide']:.3f}",
            item["driver_place"],
        ]
        for i, text in enumerate(values):
            add_cell_text(cells[i], text, size=8.5, align=WD_ALIGN_PARAGRAPH.CENTER if i in (0, 3) else None)
            set_cell_border(cells[i])

    add_heading(doc, "활용방안 및 기대효과")
    add_bullets(doc, [
        "상권·관광지 혼잡 및 활력 변화를 시간대별로 파악해 방문 분산 안내, 현장 운영 인력 배치, 주변 소상공인 프로모션 타이밍 결정에 활용할 수 있다.",
        "실시간 API와 격자 기반 예측을 결합하므로 신규 행사·날씨 변화·교통 변화가 발생했을 때 빠르게 상권 활력 변화를 재계산할 수 있다.",
        "최종 점수와 산출 방식(score_source)을 함께 제공해 정책 담당자가 직접 관측 구역과 전파 추정 구역을 구분하여 해석할 수 있다.",
    ])

    add_heading(doc, "분석툴 및 참고문헌")
    add_bullets(doc, [
        "Python, pandas, NumPy, scikit-learn, PyTorch, SciPy, Matplotlib",
        "모델 구성: Base MLP, Correction LSTM, HuberLoss, AdamW, ReduceLROnPlateau, Early Stopping, Gradient Clipping",
        "활용 파일: seoul_grid_vitality_pipeline.py, preprocess_seoul_grid_data.py, fetch_seoul_realtime_api_to_csv.py, build_citywide_vitality_artifacts.py",
        "참고 데이터: 서울시 빅데이터캠퍼스 제공 데이터, 서울 실시간 도시데이터 OpenAPI, 행정구 경계 및 격자·장소 매핑 산출물",
    ])

    path = OUT / f"{TEAM_PLACEHOLDER}_분석요약서.docx"
    doc.save(path)
    return path


def make_application_doc(summary):
    doc = Document()
    set_doc_defaults(doc)
    add_title(doc, "2026 서울시 빅데이터 활용 경진대회 분석 부문", "참 가 신 청 서")

    info = doc.add_table(rows=8, cols=4)
    info.alignment = WD_TABLE_ALIGNMENT.CENTER
    fields = [
        ("팀명", TEAM_PLACEHOLDER, "개인 참가자는 성명 기재", ""),
        ("성명", "[대표자 성명]", "성별", "[성별]"),
        ("생년월일", "[YYYY.MM.DD]", "전화번호", "[전화번호]"),
        ("주소", "[주민등록상 주소]", "이메일", "[이메일]"),
        ("소속", "[직장 또는 학교명]", "팀원", "[팀원 성명/역할]"),
        ("공모전", "2026 빅데이터 활용 경진대회 분석부문", "분석명", PROJECT_TITLE),
        ("제출물", "분석결과서(PPT), 분석요약서, 소스 및 데이터", "동의", "동의함"),
        ("작성일", "2026.    .    .", "대표자", "[서명]"),
    ]
    for r, row in enumerate(fields):
        for c, text in enumerate(row):
            cell = info.cell(r, c)
            is_label = c in (0, 2)
            add_cell_text(cell, text, bold=is_label, color="FFFFFF" if is_label else "1F2937", align=WD_ALIGN_PARAGRAPH.CENTER if is_label else None)
            if is_label:
                set_cell_shading(cell, "16655C")
            set_cell_border(cell)

    add_heading(doc, "추진배경 및 필요성")
    add_bullets(doc, [
        "상권·관광지의 시간대별 활력은 인구, 소비, 날씨, 교통, 행사 요인이 동시에 작용해 빠르게 변화한다.",
        "정적 상권 지표와 실시간 도시데이터를 결합하여 현장 운영 및 공공정책 의사결정에 바로 활용할 수 있는 격자 단위 점수가 필요하다.",
    ])

    add_heading(doc, "분석내용(투입데이터)")
    add_bullets(doc, [
        "서울 50m 격자 단위로 유동인구, 카드 소비, 강우, 교통 접근성, 임대료 proxy 등 평시 상권 특성을 구성했다.",
        "서울 실시간 도시데이터 API의 인구·교통·날씨·행사 신호를 장소-grid 매핑으로 변환했다.",
        "Base MLP와 Correction LSTM을 결합해 final_score를 산출하고, 직접 관측 및 전파 추정 격자를 구분했다.",
        f"분석 범위는 {summary['covered_grid_count']}개 격자, 직접 실시간 관측 {summary['live_grid_count']}개 격자, 주요 장소 {summary['place_count']}개이다.",
    ])

    add_heading(doc, "활용방안 및 기대효과")
    add_bullets(doc, [
        "상권 활력 상위·변동 지역을 파악해 관광 동선, 혼잡 완화, 소상공인 지원, 현장 인력 배치에 활용한다.",
        "실시간 보정 체계를 통해 날씨·행사·교통 변화에 따른 상권 반응을 빠르게 재계산할 수 있다.",
    ])

    add_heading(doc, "준수사항")
    add_body(doc, "본 신청서는 제공 양식의 주요 항목을 기준으로 작성한 초안입니다. 팀명, 대표자, 연락처, 주소, 소속, 팀원 정보 및 서명은 제출 전 참가자가 직접 확인·기입해야 합니다.")
    add_body(doc, "대회 제반 규정을 준수하며, 제출 산출물의 권리·공개·교육 활용 관련 동의 여부는 제출 전 최종 확인합니다.")

    path = OUT / f"{TEAM_PLACEHOLDER}_참가신청서_작성본.docx"
    doc.save(path)
    return path


if __name__ == "__main__":
    summary = load_summary()
    print(make_summary_doc(summary))
    print(make_application_doc(summary))
